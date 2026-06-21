import os, docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def set_run_fonts(run, ascii_font, ascii_size, east_font, east_size):
    rPr = run._element.find('w:rPr', ns)
    if rPr is None:
        rPr = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', {})
        run._element.insert(0, rPr)
    for child in list(rPr):
        tag = child.tag.split('}')[-1]
        if tag == 'rFonts':
            rPr.remove(child)
    rFonts = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts', {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia': east_font,
    })
    rPr.insert(0, rFonts)
    sz = rPr.find('w:sz', ns)
    if sz is None:
        sz = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', {})
        rPr.append(sz)
    sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(ascii_size.pt * 2)))


def split_run_mixed(run):
    text = run.text
    if text == '':
        return [run]
    spans = []
    cur_cls = None; start = 0
    for i, ch in enumerate(text):
        c = 'A' if ord(ch) < 128 else 'N'
        if cur_cls is None:
            cur_cls = c; start = i
        elif c != cur_cls:
            spans.append((cur_cls, start, i))
            cur_cls = c; start = i
    spans.append((cur_cls, start, len(text)))
    if len(spans) == 1:
        return [run]
    parent = run._element.getparent()
    idx = parent.index(run._element)
    new_runs = []
    for cls, s, e in spans:
        r_elem = run._element.makeelement(run._element.tag, run._element.attrib)
        t_elem = r_elem.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', {})
        t_elem.text = text[s:e]
        r_elem.append(t_elem)
        parent.insert(idx, r_elem)
        idx += 1
        new_runs.append(docx.text.run.Run(r_elem, run._parent))
    parent.remove(run._element)
    return new_runs


def fix_docx(path, chinese):
    doc = docx.Document(path)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    try:
        normal_style = doc.styles['Normal']
        if chinese:
            normal_style.font.name = 'Microsoft YaHei'
    except KeyError:
        pass
    fixed_code = 0
    for para in doc.paragraphs:
        style = (para.style.name if para.style else '').lower()
        is_code = ('code' in style or 'source' in style or 'verbatim' in style)
        if is_code:
            pPr = para._element.find('w:pPr', ns)
            if pPr is None:
                pPr = para._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr', {})
                para._element.insert(0, pPr)
            wrap = pPr.find('w:wrap', ns)
            if wrap is None:
                wrap = pPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}wrap', {
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'linesAndChars'
                })
                pPr.append(wrap)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fixed_code += 1
        new_runs = []
        for run in list(para.runs):
            new_runs.extend(split_run_mixed(run))
        for r in new_runs:
            has_cjk = any(ord(c) >= 128 for c in r.text)
            has_ascii = any(ord(c) < 128 for c in r.text)
            if is_code:
                if chinese:
                    if has_cjk and not has_ascii:
                        set_run_fonts(r, 'Microsoft YaHei', Pt(9), 'Microsoft YaHei', Pt(9))
                    else:
                        set_run_fonts(r, 'Consolas', Pt(9), 'Microsoft YaHei', Pt(9))
                else:
                    set_run_fonts(r, 'Consolas', Pt(9), 'Calibri', Pt(10))
            else:
                if chinese and has_cjk:
                    set_run_fonts(r, 'Calibri', Pt(11), 'Microsoft YaHei', Pt(11))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        for r2 in split_run_mixed(run):
                            has_cjk = any(ord(c) >= 128 for c in r2.text)
                            if chinese and has_cjk:
                                set_run_fonts(r2, 'Calibri', Pt(10), 'Microsoft YaHei', Pt(10))
    doc.save(path)
    print(f'{os.path.basename(path)}: code_paras={fixed_code}')


root = r'd:\project\code\maths\黎曼猜想'
fix_docx(os.path.join(root, 'riemann_thesis_cn.docx'), chinese=True)
fix_docx(os.path.join(root, 'riemann_thesis_en.docx'), chinese=False)

# sanity-check chinese docx code paras: no more '?' + run fonts
d = docx.Document(os.path.join(root, 'riemann_thesis_cn.docx'))
print('--- cn docx sanity ---')
shown = 0
for para in d.paragraphs:
    s = (para.style.name if para.style else '').lower()
    if 'code' in s or 'source' in s:
        t = para.text[:80]
        has_bad = '?' in t and not ('Coq? 9.0' in t)
        print(f'[{shown}] {t!r} has_cjk={any(ord(c)>=128 for c in para.text)}')
        for r in para.runs[:3]:
            has_cjk_r = any(ord(c)>=128 for c in r.text)
            rPr = r._element.find('w:rPr', ns)
            rFonts = None
            if rPr is not None:
                for c in rPr:
                    if c.tag.endswith('rFonts'):
                        rFonts = c; break
            ascii_f = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii') if rFonts is not None else None
            east_f = rFonts.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia') if rFonts is not None else None
            print(f'    run(has_cjk={has_cjk_r}) ascii={ascii_f} eastAsia={east_f} text={r.text[:30]!r}')
        shown += 1
        if shown >= 3: break
print('done')
