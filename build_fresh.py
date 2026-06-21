import re, os, subprocess, shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = r'd:\project\code\maths\黎曼猜想'

def fix_md(src_name, dst_name):
    p = os.path.join(root, src_name)
    with open(p, 'rb') as f: b = f.read()
    while b.startswith(b'\xef\xbb\xbf'): b = b[3:]
    t = b.decode('utf-8')
    t = t.replace('✅', '[OK]')
    t = re.sub(r'\\\(\s*([\s\S]*?)\\\)', lambda m: '$' + m.group(1).strip() + '$', t)
    t = re.sub(r'\\\[\s*([\s\S]*?)\\\]', lambda m: '$$\n' + m.group(1).strip() + '\n$$', t)
    t = re.sub(r'\n> \[\s*\n', '\n> $$\n', t)
    t = re.sub(r'\n> \]\s*\n', '\n> \n', t)
    with open(os.path.join(root, dst_name), 'w', encoding='utf-8') as f: f.write(t)
    real_cjk = sum(1 for c in t if '\u4e00' <= c <= '\u9fff')
    print(f'cleaned {src_name} -> {dst_name} chars={len(t)} real_cjk={real_cjk} has_三类={("三类" in t)}')

fix_md('riemann_thesis_backup_before_Replace3.md', 'cn_fresh_ok.md')
fix_md('riemann_thesis_backup_before_Replace3.md', 'cn_fresh_ok2.md')

# pandoc via stdin -> html5 -> Word -> postprocess -> verify Chinese is real
def pandoc_stdin(md_text, out_html):
    p = subprocess.Popen(
        ['pandoc','--standalone','--from','markdown','--to','html5','--mathml','-o',out_html],
        stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        cwd=root,
        env={**os.environ, 'PYTHONUTF8':'1', 'PYTHONIOENCODING':'utf-8', 'PYTHONLEGACYWINDOWSSTDIO':'1'},
    )
    out, err = p.communicate(input=md_text.encode('utf-8'))
    if p.returncode != 0:
        print('pandoc stderr:', err.decode('utf-8','replace'))
        raise SystemExit(p.returncode)

def word_html2docx(html, docx):
    import win32com.client
    w = win32com.client.Dispatch('Word.Application')
    w.Visible = False
    d = w.Documents.Open(html, ReadOnly=True)
    d.SaveAs(docx, FileFormat=16)
    d.Close(False)
    try: w.Quit()
    except Exception: pass

def set_run_fonts(run, ascii_font, ascii_size, east_font, east_size):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    rPr = run._element.find('w:rPr', ns)
    if rPr is None:
        rPr = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', {})
        run._element.insert(0, rPr)
    for c in list(rPr):
        if c.tag.endswith('rFonts'): rPr.remove(c)
    rf = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts', {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia': east_font,
    })
    rPr.insert(0, rf)
    sz = rPr.find('w:sz', ns)
    if sz is None:
        sz = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', {})
        rPr.append(sz)
    sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(ascii_size.pt*2)))
    cs = rPr.find('w:szCs', ns)
    if cs is None:
        cs = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs', {})
        rPr.append(cs)
    cs.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(east_size.pt*2)))

def split_run(run):
    t = run.text
    if t == '': return [run]
    spans = []; cur=None; start=0
    for i, ch in enumerate(t):
        c = 'A' if ord(ch) < 128 else 'N'
        if cur is None: cur=c; start=i
        elif c != cur: spans.append((cur,start,i)); cur=c; start=i
    spans.append((cur,start,len(t)))
    if len(spans) == 1: return [run]
    parent = run._element.getparent()
    idx = parent.index(run._element)
    new_runs = []
    from docx.text.run import Run
    for cls, s, e in spans:
        r_elem = run._element.makeelement(run._element.tag, run._element.attrib)
        t_elem = r_elem.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', {})
        t_elem.text = t[s:e]
        r_elem.append(t_elem)
        parent.insert(idx, r_elem); idx += 1
        new_runs.append(Run(r_elem, run._parent))
    parent.remove(run._element)
    return new_runs

def postprocess(path, chinese):
    doc = Document(path)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for para in doc.paragraphs:
        style = (para.style.name if para.style else '').lower()
        is_code = ('code' in style or 'source' in style or 'verbatim' in style or 'pre' in style)
        if is_code:
            pPr = para._element.find('w:pPr', ns)
            if pPr is None:
                pPr = para._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr', {})
                para._element.insert(0, pPr)
            wrap = pPr.find('w:wrap', ns)
            if wrap is None:
                wrap = pPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}wrap', {
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'linesAndChars'
                })
                pPr.append(wrap)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in list(para.runs):
            for r2 in split_run(r):
                has_cjk = any(ord(c)>=128 for c in r2.text)
                has_ascii = any(ord(c)<128 for c in r2.text)
                if is_code:
                    if chinese:
                        if has_cjk and not has_ascii:
                            set_run_fonts(r2, 'Microsoft YaHei', Pt(9), 'Microsoft YaHei', Pt(9))
                        else:
                            set_run_fonts(r2, 'Consolas', Pt(9), 'Microsoft YaHei', Pt(9))
                    else:
                        set_run_fonts(r2, 'Consolas', Pt(9), 'Calibri', Pt(10))
                else:
                    if chinese and has_cjk:
                        set_run_fonts(r2, 'Calibri', Pt(11), 'Microsoft YaHei', Pt(11))
                    else:
                        set_run_fonts(r2, 'Calibri', Pt(11), 'Calibri', Pt(11))
    doc.save(path)

# Build Chinese
cn_md_p = os.path.join(root, 'cn_fresh_ok.md')
cn_text = open(cn_md_p, 'r', encoding='utf-8').read()
cn_html = os.path.join(root, 'riemann_thesis_cn.html')
cn_docx = os.path.join(root, 'riemann_thesis_cn.docx')
pandoc_stdin(cn_text, cn_html)
word_html2docx(cn_html, cn_docx)
postprocess(cn_docx, True)

# Build English
with open(os.path.join(root, 'riemann_thesis_en_clean.md'), 'r', encoding='utf-8') as f:
    en_text = f.read()
pandoc_stdin(en_text, os.path.join(root, 'riemann_thesis_en.html'))
word_html2docx(os.path.join(root, 'riemann_thesis_en.html'), os.path.join(root, 'riemann_thesis_en.docx'))
postprocess(os.path.join(root, 'riemann_thesis_en.docx'), False)

# Verify Chinese Layer0 line
d = Document(cn_docx)
ok = False
for p in d.paragraphs:
    style = (p.style.name if p.style else '').lower()
    if 'code' in style or 'source' in style or 'pre' in style:
        t = p.text
        if 'Layer 0' in t:
            first = t.split('\n')[0]
            print('Chinese Layer0:', first[:120])
            ok = ('三类' in first and '标注' in first and '变分' in t and '纯泛型' in t)
            print('  real Chinese:', ok, '  bad 涓?', '涓' in t)
            break
print('PASS' if ok else 'FAIL')
