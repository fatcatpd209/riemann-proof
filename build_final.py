import os, subprocess, shutil, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = r'd:\project\code\maths\黎曼猜想'

def read_md_utf8(p):
    with open(p,'rb') as f: b = f.read()
    if b.startswith(b'\xef\xbb\xbf'): b = b[3:]
    return b.decode('utf-8')

def pandoc_via_stdin(md_text, out_html):
    # Pass UTF-8 bytes explicitly via Popen stdin, avoid subprocess default locale encoding
    p = subprocess.Popen(
        ['pandoc','--standalone','--from','markdown','--to','html5','--mathml','-o',out_html],
        stdin=subprocess.PIPE, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        cwd=root,
        env={**os.environ, 'PYTHONUTF8':'1', 'PYTHONIOENCODING':'utf-8'},
    )
    stdout, stderr = p.communicate(input=md_text.encode('utf-8'))
    if p.returncode != 0:
        print('pandoc stderr:', stderr.decode('utf-8','replace'))
        raise SystemExit(p.returncode)
    print(f'  pandoc wrote {out_html} ({os.path.getsize(out_html)} bytes)')

def word_html_to_docx(html_path, docx_path):
    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(html_path, ReadOnly=True)
    doc.SaveAs(docx_path, FileFormat=16)  # wdFormatXMLDocument
    doc.Close()
    word.Quit()

def set_run_fonts(run, ascii_font, ascii_size, east_font, east_size):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    rPr = run._element.find('w:rPr', ns)
    if rPr is None:
        rPr = run._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr', {})
        run._element.insert(0, rPr)
    for child in list(rPr):
        if child.tag.endswith('rFonts'):
            rPr.remove(child)
    rFonts = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts', {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi': ascii_font,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia': east_font,
    })
    rPr.insert(0, rFonts)
    sz = rPr.find('w:sz', ns)
    if sz is None:
        sz = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', {})
        rPr.append(sz)
    sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(ascii_size.pt*2)))
    szE = rPr.find('w:szCs', ns)
    if szE is None:
        szE = rPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs', {})
        rPr.append(szE)
    szE.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(east_size.pt*2)))

def split_run(run):
    text = run.text
    if text == '':
        return [run]
    spans = []
    cur = None; start = 0
    for i, ch in enumerate(text):
        c = 'A' if ord(ch) < 128 else 'N'
        if cur is None:
            cur = c; start = i
        elif c != cur:
            spans.append((cur, start, i)); cur = c; start = i
    spans.append((cur, start, len(text)))
    if len(spans) == 1:
        return [run]
    parent = run._element.getparent()
    idx = parent.index(run._element)
    new_runs = []
    for cls, s, e in spans:
        r_elem = run._element.makeelement(run._element.tag, run._element.attrib)
        t_elem = r_elem.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t', {})
        t_elem.text = text[s:e]
        r_elem.append(t_elem)
        parent.insert(idx, r_elem); idx += 1
        from docx.text.run import Run
        new_runs.append(Run(r_elem, run._parent))
    parent.remove(run._element)
    return new_runs

def postprocess_docx(path, chinese):
    doc = Document(path)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    code_paras = 0
    for para in doc.paragraphs:
        style = (para.style.name if para.style else '').lower()
        is_code = ('code' in style or 'source' in style or 'verbatim' in style or 'pre' in style)
        if is_code:
            pPr = para._element.find('w:pPr', ns)
            if pPr is None:
                pPr = para._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr', {})
                para._element.insert(0, pPr)
            wrap = pPr.find('w:wrap', ns)
            if wrap is None:
                wrap = pPr.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}wrap', {
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'linesAndChars'
                })
                pPr.append(wrap)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            code_paras += 1
        new_runs = []
        for run in list(para.runs):
            new_runs.extend(split_run(run))
        for r in new_runs:
            has_cjk = any(ord(c)>=128 for c in r.text)
            has_ascii = any(ord(c)<128 for c in r.text)
            if is_code:
                if chinese:
                    if has_cjk and not has_ascii:
                        set_run_fonts(r, 'Microsoft YaHei', Pt(9), 'Microsoft YaHei', Pt(9))
                    else:
                        set_run_fonts(r, 'Consolas', Pt(9), 'Microsoft YaHei', Pt(9))
                else:
                    set_run_fonts(r, 'Consolas', Pt(9), 'Calibri', Pt(10))
            else:
                if chinese and has_cjk:
                    set_run_fonts(r, 'Calibri', Pt(11), 'Microsoft YaHei', Pt(11))
                else:
                    set_run_fonts(r, 'Calibri', Pt(11), 'Calibri', Pt(11))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        for r2 in split_run(run):
                            has_cjk = any(ord(c)>=128 for c in r2.text)
                            if chinese and has_cjk:
                                set_run_fonts(r2, 'Calibri', Pt(10), 'Microsoft YaHei', Pt(10))
                            else:
                                set_run_fonts(r2, 'Calibri', Pt(10), 'Calibri', Pt(10))
    doc.save(path)
    print(f'{os.path.basename(path)}: code_paras={code_paras}, margins 1.8cm, wrap=linesAndChars, fonts OK')

for (md, html, docx, chinese) in [
    ('riemann_thesis_cn_clean.md','riemann_thesis_cn.html','riemann_thesis_cn.docx',True),
    ('riemann_thesis_en_clean.md','riemann_thesis_en.html','riemann_thesis_en.docx',False),
]:
    print(f'--- building {md} ---')
    mdt = read_md_utf8(os.path.join(root, md))
    pandoc_via_stdin(mdt, os.path.join(root, html))
    word_html_to_docx(os.path.join(root, html), os.path.join(root, docx))
    postprocess_docx(os.path.join(root, docx), chinese)

# verify Chinese code block text is real Chinese
d = Document(os.path.join(root, 'riemann_thesis_cn.docx'))
print('--- cn docx code block sanity ---')
for p in d.paragraphs:
    s = (p.style.name if p.style else '').lower()
    if 'code' in s or 'source' in s or 'pre' in s:
        t = p.text
        if 'Layer 0' in t:
            line = t.split('\n')[0]
            print('Layer0 line:', line)
            print('  has real Chinese 三类:', '三类' in line)
            print('  has real Chinese 标注:', '标注' in line)
            print('  has real Chinese 纯泛型:', '纯泛型' in t)
            print('  has bad mojibake 涓:', '涓' in t)
            break
