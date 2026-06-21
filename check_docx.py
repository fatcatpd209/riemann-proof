import os, sys, zipfile, re
try:
    import docx
except ImportError:
    os.system('python -m pip install python-docx --quiet')
    import docx

root = r'd:\project\code\maths\黎曼猜想'

def analyze(path):
    d = docx.Document(path)
    print(f'=== {os.path.basename(path)} ===')
    print(f'paragraphs: {len(d.paragraphs)}, tables: {len(d.tables)}')
    # collect
    code_paras = []  # indices of plaintext paras that look like code
    para_info = []
    for i,p in enumerate(d.paragraphs):
        t = p.text
        style = p.style.name if p.style else ''
        # detect code blocks: pandoc usually uses "Quote" + run font="Courier New" or similar
        runs_info = []
        for r in p.runs:
            f = r.font
            runs_info.append({
                'text': r.text[:40],
                'font': f.name,
                'size': None if f.size is None else f.size.pt,
                'italic': f.italic,
                'bold': f.bold,
            })
        is_code = (
            ('Code' in style or 'Quote' in style or 'Verbatim' in style)
            or any((r.get('font') and r['font'].lower().startswith(('courier','consolas','mono','source','dejavu sans mono')) for r in runs_info))
            or t.startswith(('def ','import ','class ','#include','using ','int ','bool ','lemma ','theorem','Fixpoint ','Inductive ','Lemma ','Theorem '))
        )
        # long line detection
        lines = t.split('\n')
        maxlen = max((len(l) for l in lines), default=0)
        para_info.append({
            'idx': i,
            'style': style,
            'text_preview': t[:80],
            'n_chars': len(t),
            'n_lines': len(lines),
            'max_line_len': maxlen,
            'is_code': is_code and len(t.strip()) > 10,
            'runs': runs_info[:3],
        })
        if is_code and len(t.strip()) > 10:
            code_paras.append((i, maxlen, t[:60].replace('\n','\\n')))
    print('  code-ish paragraphs:', len(code_paras))
    # show worst 10 longest lines across ALL paragraphs
    longest = sorted(para_info, key=lambda x: x['max_line_len'], reverse=True)[:10]
    print('  top-10 longest-line paragraphs:')
    for x in longest:
        flag = '[CODE]' if x['is_code'] else ''
        print(f"    #{x['idx']:>3} style={x['style']:<12} maxlen={x['max_line_len']:>5} {flag} {x['text_preview']}")

    # tables: count cells, max text length per cell
    print('  table cells analysis:')
    for ti,t in enumerate(d.tables):
        rows = len(t.rows); cols = len(t.columns)
        max_cell = 0; max_cell_text = ''
        for r in t.rows:
            for c in r.cells:
                n = len(c.text.strip())
                if n > max_cell:
                    max_cell = n; max_cell_text = c.text.strip()[:60].replace('\n','\\n')
        print(f"    table#{ti} rows={rows} cols={cols} max_cell_chars={max_cell} -> {max_cell_text}")

    # images: media files count
    with zipfile.ZipFile(path,'r') as z:
        media = [n for n in z.namelist() if n.startswith('word/media/')]
    print(f'  embedded images/media: {len(media)}')

    # suspicious chars check (non-UTF8 replacement char or common mojibake)
    bad = []
    for i,p in enumerate(d.paragraphs):
        if '\ufffd' in p.text:
            bad.append((i,p.text))
    if bad:
        print('  WARNING: replacement chars:', len(bad))
        for i,t in bad[:5]:
            print(f'    para#{i}: {t[:100]}')
    else:
        print('  no replacement chars in paragraphs')

for n in ('riemann_thesis_cn.docx','riemann_thesis_en.docx'):
    analyze(os.path.join(root,n))
    print()
